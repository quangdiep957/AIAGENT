"""
File Reader Tool - Đọc nội dung từ các loại file khác nhau
Hỗ trợ: PDF, Word (DOC/DOCX), Text files
"""

import os
import re
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime

# Import các thư viện xử lý file
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

class FileReaderTool:
    """Tool đọc nội dung từ các loại file"""
    
    def __init__(self):
        """Khởi tạo FileReaderTool"""
        self.supported_formats = {
            'pdf': self._read_pdf if PDF_AVAILABLE else None,
            'docx': self._read_docx if DOCX_AVAILABLE else None,
            'doc': self._read_docx if DOCX_AVAILABLE else None,
            'txt': self._read_text,
            'md': self._read_text,
            'rtf': self._read_text
        }
    
    def _detect_encoding(self, file_path: str) -> str:
        """
        Tự động phát hiện encoding của file text
        
        Args:
            file_path (str): Đường dẫn file
            
        Returns:
            str: Encoding được phát hiện
        """
        if CHARDET_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                    result = chardet.detect(raw_data)
                    return result.get('encoding', 'utf-8')
            except:
                pass
        
        # Fallback: thử các encoding phổ biến
        encodings = ['utf-8', 'utf-16', 'windows-1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    f.read(100)  # Đọc thử 100 ký tự đầu
                return encoding
            except:
                continue
        
        return 'utf-8'  # Default
    
    def _clean_text(self, text: str) -> str:
        """
        Làm sạch text: loại bỏ ký tự không cần thiết
        
        Args:
            text (str): Text cần làm sạch
            
        Returns:
            str: Text đã được làm sạch
        """
        if not text:
            return ""
        
        # Loại bỏ ký tự điều khiển
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Chuẩn hóa khoảng trắng
        text = re.sub(r'\s+', ' ', text)
        
        # Loại bỏ khoảng trắng đầu/cuối
        text = text.strip()
        
        return text
    
    def _read_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Đọc nội dung từ file PDF
        
        Args:
            file_path (str): Đường dẫn file PDF
            
        Returns:
            Dict[str, Any]: Nội dung đã được extract
        """
        try:
            pages_content = []
            total_text = ""
            
            # Thử dùng pdfplumber trước (tốt hơn cho tables)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text() or ""
                        page_text = self._clean_text(page_text)
                        
                        pages_content.append({
                            "page_number": page_num,
                            "content": page_text,
                            "word_count": len(page_text.split()) if page_text else 0
                        })
                        
                        total_text += page_text + "\n"
                        
            except Exception as e:
                # Fallback sang PyPDF2
                print(f"pdfplumber failed, trying PyPDF2: {e}")
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    pages_content = []
                    total_text = ""
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text() or ""
                        page_text = self._clean_text(page_text)
                        
                        pages_content.append({
                            "page_number": page_num,
                            "content": page_text,
                            "word_count": len(page_text.split()) if page_text else 0
                        })
                        
                        total_text += page_text + "\n"
            
            return {
                "success": True,
                "file_type": "pdf",
                "pages": pages_content,
                "total_pages": len(pages_content),
                "total_content": self._clean_text(total_text),
                "total_word_count": len(total_text.split()) if total_text else 0,
                "extracted_date": datetime.utcnow()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Lỗi khi đọc PDF: {str(e)}",
                "file_type": "pdf"
            }
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Đọc nội dung từ file Word (DOCX)
        
        Args:
            file_path (str): Đường dẫn file Word
            
        Returns:
            Dict[str, Any]: Nội dung đã được extract
        """
        try:
            doc = Document(file_path)
            
            paragraphs = []
            tables_content = []
            total_text = ""
            
            # Đọc paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                if para_text:  # Chỉ lưu paragraph không rỗng
                    para_text = self._clean_text(para_text)
                    paragraphs.append({
                        "paragraph_number": para_num,
                        "content": para_text,
                        "word_count": len(para_text.split()) if para_text else 0
                    })
                    total_text += para_text + "\n"
            
            # Đọc tables
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = self._clean_text(cell.text)
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                if table_data:  # Chỉ lưu table không rỗng
                    tables_content.append({
                        "table_number": table_num,
                        "data": table_data,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0
                    })
                    
                    # Thêm nội dung table vào total_text
                    for row in table_data:
                        total_text += " | ".join(row) + "\n"
            
            return {
                "success": True,
                "file_type": "docx",
                "paragraphs": paragraphs,
                "tables": tables_content,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables_content),
                "total_content": self._clean_text(total_text),
                "total_word_count": len(total_text.split()) if total_text else 0,
                "extracted_date": datetime.utcnow()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Lỗi khi đọc Word document: {str(e)}",
                "file_type": "docx"
            }
    
    def _read_text(self, file_path: str) -> Dict[str, Any]:
        """
        Đọc nội dung từ file text (TXT, MD, RTF)
        
        Args:
            file_path (str): Đường dẫn file text
            
        Returns:
            Dict[str, Any]: Nội dung đã được extract
        """
        try:
            # Tự động phát hiện encoding
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            # Làm sạch content
            content = self._clean_text(content)
            
            # Chia thành lines
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            
            return {
                "success": True,
                "file_type": "text",
                "content": content,
                "lines": lines,
                "total_lines": len(lines),
                "total_word_count": len(content.split()) if content else 0,
                "encoding_used": encoding,
                "extracted_date": datetime.utcnow()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Lỗi khi đọc text file: {str(e)}",
                "file_type": "text"
            }
    
    def read_file(self, file_path: str) -> Dict[str, Any]:
        """
        Đọc nội dung từ file (tự động detect loại file)
        
        Args:
            file_path (str): Đường dẫn file
            
        Returns:
            Dict[str, Any]: Nội dung và metadata của file
        """
        try:
            # Kiểm tra file có tồn tại
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": "File không tồn tại"
                }
            
            # Xác định extension
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            
            # Kiểm tra format có được hỗ trợ không
            if file_extension not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Định dạng file '{file_extension}' không được hỗ trợ. "
                            f"Các định dạng hỗ trợ: {list(self.supported_formats.keys())}"
                }
            
            # Kiểm tra thư viện có sẵn không
            reader_func = self.supported_formats[file_extension]
            if reader_func is None:
                return {
                    "success": False,
                    "error": f"Thư viện xử lý file '{file_extension}' chưa được cài đặt"
                }
            
            # Đọc file
            result = reader_func(file_path)
            
            # Thêm metadata chung
            if result.get("success"):
                file_stat = os.stat(file_path)
                result.update({
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_size": file_stat.st_size,
                    "file_extension": file_extension,
                    "processing_date": datetime.utcnow()
                })
            
            return result
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Lỗi khi đọc file: {str(e)}"
            }
    
    def extract_chunks(self, content: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """
        Chia nội dung thành các chunks nhỏ để xử lý embedding
        
        Args:
            content (str): Nội dung cần chia
            chunk_size (int): Kích thước chunk (số ký tự)
            overlap (int): Số ký tự overlap giữa các chunk
            
        Returns:
            List[Dict]: Danh sách các chunks
        """
        if not content:
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(content):
            # Xác định điểm kết thúc chunk
            end = start + chunk_size
            
            # Nếu không phải chunk cuối, tìm điểm cắt hợp lý (sau dấu câu)
            if end < len(content):
                # Tìm dấu câu gần nhất trong khoảng overlap
                for i in range(end, max(start + chunk_size - overlap, end - 100), -1):
                    if content[i] in '.!?;:\n':
                        end = i + 1
                        break
            
            # Extract chunk
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                chunks.append({
                    "chunk_index": chunk_index,
                    "content": chunk_content,
                    "start_position": start,
                    "end_position": end,
                    "word_count": len(chunk_content.split()),
                    "char_count": len(chunk_content)
                })
                chunk_index += 1
            
            # Di chuyển start position với overlap
            start = end - overlap if end < len(content) else end
        
        return chunks
    
    def get_reading_time(self, word_count: int, wpm: int = 200) -> Dict[str, Any]:
        """
        Ước tính thời gian đọc
        
        Args:
            word_count (int): Số từ
            wpm (int): Words per minute (mặc định 200)
            
        Returns:
            Dict[str, Any]: Thông tin thời gian đọc
        """
        if word_count <= 0:
            return {"minutes": 0, "seconds": 0, "formatted": "0 phút"}
        
        total_minutes = word_count / wpm
        minutes = int(total_minutes)
        seconds = int((total_minutes - minutes) * 60)
        
        if minutes > 0:
            formatted = f"{minutes} phút"
            if seconds > 0:
                formatted += f" {seconds} giây"
        else:
            formatted = f"{seconds} giây"
        
        return {
            "minutes": minutes,
            "seconds": seconds,
            "total_minutes": total_minutes,
            "formatted": formatted,
            "wpm_used": wpm
        }

# Example usage
if __name__ == "__main__":
    # Demo usage
    reader = FileReaderTool()
    
    print("=== File Reader Tool Demo ===")
    
    # Tạo file test
    test_file = "test_content.txt"
    with open(test_file, "w", encoding="utf-8") as f:
        f.write("""
        Đây là nội dung test cho File Reader Tool.
        
        Tool này có thể đọc nhiều loại file khác nhau:
        - File PDF
        - File Word (DOC/DOCX)
        - File Text (TXT, MD, RTF)
        
        Nội dung sẽ được làm sạch và chuẩn hóa.
        Tool cũng hỗ trợ chia nội dung thành chunks để xử lý embedding.
        """)
    
    # Đọc file
    result = reader.read_file(test_file)
    print(f"Reading result: {result}")
    
    if result.get("success"):
        # Chia thành chunks
        content = result.get("content", "")
        chunks = reader.extract_chunks(content, chunk_size=100, overlap=20)
        print(f"\nChunks created: {len(chunks)}")
        for chunk in chunks:
            print(f"Chunk {chunk['chunk_index']}: {chunk['content'][:50]}...")
        
        # Ước tính thời gian đọc
        word_count = result.get("total_word_count", 0)
        reading_time = reader.get_reading_time(word_count)
        print(f"\nReading time: {reading_time['formatted']}")
    
    # Cleanup
    if os.path.exists(test_file):
        os.remove(test_file)
